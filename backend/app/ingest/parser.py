"""文档解析：Markdown / TXT / PDF / DOCX → 纯文本 + 标题元信息。

按 source_format 分派：
- md   → parse_markdown (正则提取 # 标题层级)
- txt  → parse_text (直接返回纯文本，无标题层级)
- pdf  → parse_pdf (pypdf 逐页提取文本，无标题层级)
- docx → parse_docx (python-docx 按 heading 样式切层级)

article 类型固定走 parse_markdown（文章源码即 Markdown）。
"""
from dataclasses import dataclass


@dataclass
class ParsedDoc:
    """解析结果。"""

    title: str
    text: str
    headings: list[dict]  # [{level, text, position}]
    slug: str | None = None


def parse_markdown(title: str, content_md: str, slug: str | None = None) -> ParsedDoc:
    """解析 Markdown：直接取文本，按 # 提取标题层级。

    P0 简化：不做完整 AST，只用正则提取 # 标题作为分块边界提示。
    position 是字符偏移量（行首在 content_md 中的位置），供 chunker 按字符切片。
    """
    import re

    headings: list[dict] = []
    offset = 0
    for line in content_md.splitlines(keepends=True):
        m = re.match(r"^(#{1,6})\s+(.+?)\s*\r?\n?$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            headings.append({"level": level, "text": text, "position": offset})
        offset += len(line)

    return ParsedDoc(
        title=title,
        text=content_md,
        headings=headings,
        slug=slug,
    )


def parse_text(title: str, text: str) -> ParsedDoc:
    """解析纯文本：无标题层级，整体作为一块文本交由 chunker 滑窗分块。"""
    return ParsedDoc(title=title, text=text, headings=[], slug=None)


def parse_pdf(title: str, raw: bytes) -> ParsedDoc:
    """解析 PDF：pypdf 逐页提取文本，拼接为纯文本。

    PDF 无可靠的标题层级结构，headings 留空，chunker 退化为纯滑窗分块。
    扫描版 PDF（图片型）提取结果为空，由 pipeline 标记 failed 并提示用户。
    """
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            # 单页提取失败不阻断整体，跳过该页
            page_text = ""
        if page_text:
            parts.append(page_text)

    text = "\n\n".join(parts)
    if not text.strip():
        raise ValueError(
            "PDF 未提取到文本（可能是扫描版/图片型 PDF，暂不支持 OCR）"
        )
    return ParsedDoc(title=title, text=text, headings=[], slug=None)


def parse_docx(title: str, raw: bytes) -> ParsedDoc:
    """解析 DOCX：python-docx 按 heading 样式切层级。

    遍历段落，Heading 1-6 样式作为标题（记录字符偏移供 chunker 切边界），
    其余段落作为正文。表格单元格文本按行拼接为正文。
    """
    import io

    from docx import Document

    doc = Document(io.BytesIO(raw))

    parts: list[str] = []
    headings: list[dict] = []
    offset = 0

    def _append(text: str) -> None:
        nonlocal offset
        if not text:
            return
        # 统一换行，便于 chunker 按行处理
        line = text if text.endswith("\n") else text + "\n"
        parts.append(line)
        offset += len(line)

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        # Word 标题样式名：Heading 1 / heading 1 / 标题 1
        level = _heading_level(style)
        if level:
            headings.append({"level": level, "text": text, "position": offset})
        _append(text)

    # 表格内容追加为正文（按行拼接）
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                _append(line)

    text = "".join(parts)
    if not text.strip():
        raise ValueError("DOCX 未提取到文本（可能是空文档）")
    return ParsedDoc(title=title, text=text, headings=headings, slug=None)


def _heading_level(style_name: str) -> int | None:
    """从 Word 样式名提取标题层级，非标题返回 None。

    支持：Heading 1 / heading 1 / 标题 1 / Title（视为 1 级）
    """
    if not style_name:
        return None
    # 中文"标题 N"
    import re

    m = re.match(r"^(?:heading|标题)\s*(\d)$", style_name)
    if m:
        return int(m.group(1))
    if style_name in ("title", "标题"):
        return 1
    return None


def parse_document(
    title: str,
    source_format: str,
    *,
    text: str | None = None,
    raw: bytes | None = None,
    slug: str | None = None,
) -> ParsedDoc:
    """按 source_format 分派解析器。

    - md/txt：传 text（已解码文本）
    - pdf/docx：传 raw（原始字节）
    - article（无 source_format）：等价于 md
    """
    if source_format in (None, "md"):
        if text is None:
            raise ValueError("md/txt 解析需要 text 参数")
        return parse_markdown(title, text, slug=slug)
    if source_format == "txt":
        if text is None:
            raise ValueError("txt 解析需要 text 参数")
        return parse_text(title, text)
    if source_format == "pdf":
        if raw is None:
            raise ValueError("pdf 解析需要 raw 参数")
        return parse_pdf(title, raw)
    if source_format == "docx":
        if raw is None:
            raise ValueError("docx 解析需要 raw 参数")
        return parse_docx(title, raw)
    raise ValueError(f"不支持的 source_format: {source_format}")
